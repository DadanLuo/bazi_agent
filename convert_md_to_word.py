#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown 转 Word/PDF 转换脚本
将 LangGraph 全流程面试拷打文档.md 转换为 Word 和 PDF 格式
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def parse_markdown(content: str) -> list:
    """解析 Markdown 内容为结构化数据"""
    lines = content.split('\n')
    parsed_blocks = []
    current_block = None
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 标题处理
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            parsed_blocks.append({
                'type': 'heading',
                'level': level,
                'text': text
            })
            i += 1
            continue
        
        # 代码块处理
        if line.startswith('```'):
            code_lines = []
            code_type = ''
            # 获取代码类型
            type_match = re.match(r'^```(\w*)', line)
            if type_match:
                code_type = type_match.group(1)
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            parsed_blocks.append({
                'type': 'code',
                'language': code_type,
                'content': '\n'.join(code_lines)
            })
            i += 1
            continue
        
        # 表格处理
        if '|' in line and line.strip().startswith('|'):
            table_rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                row_content = lines[i].strip()
                # 跳过分隔行
                if not re.match(r'^\|\s*[-:]+\s*\|', row_content):
                    cells = [cell.strip() for cell in row_content.split('|')[1:-1]]
                    table_rows.append(cells)
                i += 1
            if table_rows:
                parsed_blocks.append({
                    'type': 'table',
                    'rows': table_rows
                })
            continue
        
        # 列表处理
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
        if list_match:
            list_items = []
            is_ordered = list_match.group(2).rstrip('.').isdigit()
            while i < len(lines):
                list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', lines[i])
                if list_match:
                    indent = len(list_match.group(1))
                    item_text = list_match.group(3).strip()
                    list_items.append({
                        'indent': indent,
                        'text': item_text
                    })
                    i += 1
                else:
                    break
            parsed_blocks.append({
                'type': 'list',
                'ordered': is_ordered,
                'items': list_items
            })
            continue
        
        # 水平线
        if re.match(r'^---+$', line.strip()):
            parsed_blocks.append({'type': 'separator'})
            i += 1
            continue
        
        # 普通段落
        if line.strip():
            paragraph_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not any([
                re.match(r'^#{1,6}\s+', lines[i]),
                lines[i].startswith('```'),
                '|' in lines[i] and lines[i].strip().startswith('|'),
                re.match(r'^(\s*)([-*+]|\d+\.)\s+', lines[i]),
                re.match(r'^---+$', lines[i].strip())
            ]):
                paragraph_lines.append(lines[i])
                i += 1
            parsed_blocks.append({
                'type': 'paragraph',
                'content': '\n'.join(paragraph_lines)
            })
            continue
        
        i += 1
    
    return parsed_blocks


def add_code_block(doc, code_content: str, language: str = ''):
    """添加代码块到文档"""
    # 添加代码块段落
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    
    # 添加语言标识
    if language:
        lang_run = paragraph.add_run(f'[{language}]\n')
        lang_run.font.bold = True
        lang_run.font.size = Pt(8)
        lang_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 添加代码内容
    code_run = paragraph.add_run(code_content)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)
    
    return paragraph


def add_table(doc, rows: list):
    """添加表格到文档"""
    if not rows:
        return
    
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    # 设置表头样式
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text
                # 表头加粗
                if i == 0:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True


def convert_markdown_to_word(md_path: str, output_path: str):
    """将 Markdown 文件转换为 Word 文档"""
    # 读取 Markdown 文件
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 解析 Markdown
    parsed_blocks = parse_markdown(content)
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置中文字体支持
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    # 定义标题样式
    heading_styles = {
        1: ('标题 1', Pt(22), True, RGBColor(41, 128, 185)),
        2: ('标题 2', Pt(16), True, RGBColor(39, 174, 96)),
        3: ('标题 3', Pt(14), True, RGBColor(192, 57, 43)),
        4: ('标题 4', Pt(12), True, RGBColor(142, 68, 173)),
        5: ('标题 5', Pt(11), True, RGBColor(44, 62, 80)),
        6: ('标题 6', Pt(10.5), True, RGBColor(127, 140, 141)),
    }
    
    # 处理每个块
    for block in parsed_blocks:
        block_type = block.get('type')
        
        if block_type == 'heading':
            level = block['level']
            text = block['text']
            style_name, font_size, bold, color = heading_styles.get(level, ('正文', Pt(10.5), False, RGBColor(0, 0, 0)))
            
            # 尝试使用内置样式
            try:
                heading_style = doc.styles[f'Heading {level}']
            except KeyError:
                heading_style = doc.styles.add_style(f'Custom Heading {level}', WD_STYLE_TYPE.PARAGRAPH)
                heading_style.font.name = '微软雅黑'
                heading_style.font.size = font_size
                heading_style.font.bold = bold
                heading_style.font.color.rgb = color
            
            paragraph = doc.add_heading(text, level=level if level <= 9 else 0)
            if level > 9:
                for run in paragraph.runs:
                    run.font.size = font_size
                    run.font.bold = bold
                    run.font.color.rgb = color
        
        elif block_type == 'paragraph':
            text = block['content']
            # 处理行内格式
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
            text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
            text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
            
            paragraph = doc.add_paragraph()
            # 简单处理，直接添加文本
            para = doc.add_paragraph()
            for run_text in text.split('\n'):
                run = para.add_run(run_text)
                run.font.name = '微软雅黑'
                run.font.size = Pt(10.5)
        
        elif block_type == 'code':
            add_code_block(doc, block['content'], block.get('language', ''))
        
        elif block_type == 'table':
            add_table(doc, block['rows'])
        
        elif block_type == 'list':
            is_ordered = block.get('ordered', False)
            for item in block['items']:
                indent = item.get('indent', 0)
                text = item['text']
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(indent / 2)
                
                if is_ordered:
                    para.add_run(f'• {text}')
                else:
                    para.add_run(f'• {text}')
        
        elif block_type == 'separator':
            doc.add_paragraph('─' * 50)
    
    # 保存文档
    doc.save(output_path)
    print(f'Word 文档已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    md_file = 'LangGraph 全流程面试拷打文档.md'
    word_file = 'LangGraph 全流程面试拷打文档.docx'
    
    if Path(md_file).exists():
        convert_markdown_to_word(md_file, word_file)
    else:
        print(f'错误：找不到文件 {md_file}')
